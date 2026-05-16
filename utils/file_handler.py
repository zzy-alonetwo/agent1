import os
import json
import hashlib
from typing import Optional, List
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    CSVLoader,
)
from utils.logger_handler import logger

# 延迟导入，避免启动时加载失败
docx_loader_available = False
xlsx_loader_available = False
md_loader_available = False
html_loader_available = False

try:
    from docx import Document as DocxDocument
    docx_loader_available = True
except ImportError:
    logger.warning("python-docx not installed, docx support disabled")

try:
    from openpyxl import load_workbook
    xlsx_loader_available = True
except ImportError:
    logger.warning("openpyxl not installed, xlsx support disabled")

try:
    import markdown
    md_loader_available = True
except ImportError:
    logger.warning("markdown not installed, md support disabled")

try:
    from html.parser import HTMLParser
    html_loader_available = True
except ImportError:
    logger.warning("html.parser not available, html support disabled")


def get_file_md5_hex(filepath: str):
    if not os.path.exists(filepath):
        logger.error(f"[md5计算]文件{filepath}不存在")
        return None
    if not os.path.isfile(filepath):
        logger.error(f"[md5计算]路径{filepath}不是文件")
        return None

    md5_object = hashlib.md5()
    chunk_size = 4096
    try:
        with open(filepath, "rb") as f:
            while chunk := f.read(chunk_size):
                md5_object.update(chunk)
            return md5_object.hexdigest()
    except Exception as e:
        logger.error(f"计算文件{filepath}md5失败,{str(e)}")
        return None


def listidr_with_allowed_type(path: str, allowed_type):
    files = []
    if not os.path.isdir(path):
        logger.error(f"[listidr_with_allowed_type]{path}不是文件夹")
        return []

    if isinstance(allowed_type, list):
        allowed_type = tuple(allowed_type)

    for f in os.listdir(path):
        if f.endswith(allowed_type):
            files.append(os.path.join(path, f))
    return files


def pdf_loader(filepath: str, password: Optional[str] = None) -> List[Document]:
    """加载 PDF 文件"""
    try:
        return PyPDFLoader(filepath, password=password).load()
    except Exception as e:
        logger.error(f"[pdf_loader]{filepath}加载失败: {str(e)}")
        return []


def txt_loader(filepath: str, encoding: str = "utf-8") -> List[Document]:
    """加载 TXT 文件"""
    try:
        return TextLoader(filepath, encoding=encoding).load()
    except Exception as e:
        logger.error(f"[txt_loader]{filepath}加载失败: {str(e)}")
        return []


def docx_loader(filepath: str) -> List[Document]:
    """加载 DOCX 文件（Word 2007+）"""
    if not docx_loader_available:
        logger.error(f"[docx_loader]{filepath}加载失败: python-docx not installed")
        return []
    
    try:
        doc = DocxDocument(filepath)
        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        if not content.strip():
            return []
        return [Document(page_content=content, metadata={"source": filepath})]
    except Exception as e:
        logger.error(f"[docx_loader]{filepath}加载失败: {str(e)}")
        return []


def xlsx_loader(filepath: str) -> List[Document]:
    """加载 XLSX 文件（Excel 2007+）"""
    if not xlsx_loader_available:
        logger.error(f"[xlsx_loader]{filepath}加载失败: openpyxl not installed")
        return []
    
    try:
        wb = load_workbook(filepath, read_only=True)
        content_parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.iter_rows(values_only=True):
                row_str = '\t'.join([str(cell) if cell else '' for cell in row])
                content_parts.append(row_str)
        content = '\n'.join(content_parts)
        if not content.strip():
            return []
        return [Document(page_content=content, metadata={"source": filepath})]
    except Exception as e:
        logger.error(f"[xlsx_loader]{filepath}加载失败: {str(e)}")
        return []


def csv_loader(filepath: str, encoding: str = "utf-8") -> List[Document]:
    """加载 CSV 文件"""
    try:
        return CSVLoader(filepath, encoding=encoding).load()
    except Exception as e:
        logger.error(f"[csv_loader]{filepath}加载失败: {str(e)}")
        return []


def md_loader(filepath: str, encoding: str = "utf-8") -> List[Document]:
    """加载 Markdown 文件"""
    if not md_loader_available:
        logger.error(f"[md_loader]{filepath}加载失败: markdown not installed")
        return []
    
    try:
        with open(filepath, 'r', encoding=encoding) as f:
            content = f.read()
        if not content.strip():
            return []
        # 转换为纯文本
        html = markdown.markdown(content)
        # 简单的 HTML 到文本转换
        text = ''.join([line.strip() for line in html.split('<') if '>' in line]).replace('>', '\n')
        if not text.strip():
            text = content  # 如果转换失败，使用原始内容
        return [Document(page_content=text, metadata={"source": filepath})]
    except Exception as e:
        logger.error(f"[md_loader]{filepath}加载失败: {str(e)}")
        return []


def html_loader(filepath: str, encoding: str = "utf-8") -> List[Document]:
    """加载 HTML 文件"""
    try:
        with open(filepath, 'r', encoding=encoding) as f:
            content = f.read()
        
        if not content.strip():
            return []
        
        # 简单的 HTML 到文本转换
        import re
        # 移除脚本和样式
        content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL)
        content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL)
        # 移除标签
        text = re.sub(r'<[^>]+>', '\n', content)
        # 移除多余空白
        text = ' '.join(text.split())
        
        if not text.strip():
            return []
        return [Document(page_content=text, metadata={"source": filepath})]
    except Exception as e:
        logger.error(f"[html_loader]{filepath}加载失败: {str(e)}")
        return []


def json_loader(filepath: str, encoding: str = "utf-8") -> List[Document]:
    """加载 JSON 文件，将 JSON 内容转换为文本文档"""
    try:
        with open(filepath, 'r', encoding=encoding) as f:
            data = json.load(f)

        def flatten(obj, parent_key='', sep=' '):
            items = []
            if isinstance(obj, dict):
                for k, v in obj.items():
                    new_key = f"{parent_key}{sep}{k}" if parent_key else k
                    items.extend(flatten(v, new_key, sep=sep))
            elif isinstance(obj, list):
                for i, v in enumerate(obj):
                    new_key = f"{parent_key}{sep}{i}" if parent_key else str(i)
                    items.extend(flatten(v, new_key, sep=sep))
            else:
                items.append((parent_key, obj))
            return items

        content_parts = []
        if isinstance(data, (dict, list)):
            flattened = flatten(data)
            content_parts = [f"{k}: {v}" for k, v in flattened]
        else:
            content_parts = [str(data)]

        content = '\n'.join(content_parts)
        return [Document(page_content=content, metadata={"source": filepath})]
    except Exception as e:
        logger.error(f"[json_loader]{filepath}加载失败: {str(e)}")
        return []


LOADER_MAP = {
    "pdf": pdf_loader,
    "txt": txt_loader,
    "docx": docx_loader,
    "xlsx": xlsx_loader,
    "csv": csv_loader,
    "md": md_loader,
    "html": html_loader,
    "json": json_loader,
}


def get_file_loader(filepath: str):
    """根据文件扩展名获取对应的加载器"""
    ext = filepath.lower().split('.')[-1]
    return LOADER_MAP.get(ext)

